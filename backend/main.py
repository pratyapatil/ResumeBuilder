# from fastapi import FastAPI, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import httpx
# import os
# import json
# import logging
# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from fastapi.responses import StreamingResponse
# import io
# from decouple import config

# app = FastAPI()

# # Configure logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Configure CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # In production, specify the actual origin
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://host.docker.internal:11434")
# # MODEL_NAME = os.getenv("MODEL_NAME", "llama3.2")

# OLLAMA_HOST = config("MODEL_NAME")
# MODEL_NAME = config("MODEL_HOST_NAME")

# class ResumeAnalysisRequest(BaseModel):
#     resume_data: dict
#     job_description: str

# class AnalysisResult(BaseModel):
#     ats_score: int
#     missing_skills: list[str]
#     improvements: list[str]

# @app.post("/api/analyze-resume-stream")
# async def analyze_resume_stream(request: ResumeAnalysisRequest):
#     async def event_generator():
#         try:
#             # 1. Sanitize resume data: Remove heavy base64 'photo' property before sending to LLM
#             clean_resume = dict(request.resume_data)
#             if clean_resume.get("personalInfo") and "photo" in clean_resume["personalInfo"]:
#                 clean_resume["personalInfo"] = dict(clean_resume["personalInfo"])
#                 clean_resume["personalInfo"].pop("photo", None)

#             prompt = f"""
#             You are an expert ATS (Applicant Tracking System) analyst and career coach.
#             Analyze the following resume data against the provided job description.
            
#             ### Job Description:
#             {request.job_description}
            
#             ### Resume Data:
#             {json.dumps(clean_resume, indent=2)}
            
#             ### Goal:
#             1. ATS match score (0-100).
#             2. List missing skills from JD.
#             3. 3-5 bullet point improvements.
            
#             ### Output Format (JSON):
#             {{
#                 "ats_score": number,
#                 "missing_skills": ["skill1", "skill2"],
#                 "improvements": ["tip1", "tip2"]
#             }}
#             """

#             async with httpx.AsyncClient(timeout=600.0) as client:
#                 async with client.stream(
#                     "POST",
#                     # f"{OLLAMA_HOST}/api/generate",
#                     f"{OLLAMA_HOST}",
#                     json={
#                         "model": MODEL_NAME,
#                         "messages": [{ "role": "user", "content": prompt }],
#                         "temperature": 0.6,
#                         # "prompt": prompt,
#                         "stream": True,
#                         "format": "json",
#                         "options": {
#                             "num_ctx": 2048,       # 2. Limit context size to speed up processing
#                             "num_predict": 800     # 3. Limit max output tokens
#                         }
#                     }
#                 ) as response:
#                     full_response = ""
#                     async for line in response.aiter_lines():
#                         if not line:
#                             continue
#                         try:
#                             data = json.loads(line)
#                             chunk = data.get("response", "")
#                             full_response += chunk
#                             if chunk:
#                                 yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                            
#                             if data.get("done"):
#                                 # Final structured data
#                                 yield f"data: {json.dumps({'done': True, 'full_response': full_response})}\n\n"
#                                 break
#                         except json.JSONDecodeError:
#                             continue

#         except Exception as e:
#             logger.error(f"Streaming error: {str(e)}")
#             yield f"data: {json.dumps({'error': str(e)})}\n\n"

#     return StreamingResponse(event_generator(), media_type="text/event-stream")

# @app.post("/api/export-word")
# async def export_word(resume_data: dict):
#     try:
#         doc = Document()
        
#         # Helper to add section titles
#         def add_section_title(text):
#             p = doc.add_heading(text, level=1)
#             run = p.runs[0]
#             run.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
        
#         # Personal Info
#         pi = resume_data.get("personalInfo", {})
#         name = f"{pi.get('firstName', '')} {pi.get('lastName', '')}".strip()
#         if name:
#             h = doc.add_heading(name, 0)
#             h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
#         contact = []
#         if pi.get("email"): contact.append(pi["email"])
#         if pi.get("phone"): contact.append(pi["phone"])
#         if pi.get("address"): contact.append(pi["address"])
#         if contact:
#             cp = doc.add_paragraph(" | ".join(contact))
#             cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
#         # Summary
#         if pi.get("summary"):
#             add_section_title("Professional Summary")
#             doc.add_paragraph(pi["summary"])
        
#         # Experience
#         exp = resume_data.get("experience", [])
#         if exp:
#             add_section_title("Work Experience")
#             for job in exp:
#                 p = doc.add_paragraph()
#                 r = p.add_run(f"{job.get('company', '')} - {job.get('role', '')}")
#                 r.bold = True
#                 doc.add_paragraph(f"{job.get('startDate', '')} - {job.get('endDate', '')}", style='Italic')
#                 if job.get("description"):
#                     doc.add_paragraph(job["description"])
        
#         # Education
#         edu = resume_data.get("education", [])
#         if edu:
#             add_section_title("Education")
#             for school in edu:
#                 p = doc.add_paragraph()
#                 r = p.add_run(f"{school.get('school', '')} - {school.get('degree', '')}")
#                 r.bold = True
#                 doc.add_paragraph(f"{school.get('startDate', '')} - {school.get('endDate', '')}", style='Italic')
        
#         # Skills
#         skills = resume_data.get("skills", [])
#         if skills:
#             add_section_title("Skills")
#             skill_list = [s.get("name", "") for s in skills if s.get("name")]
#             doc.add_paragraph(", ".join(skill_list))
            
#         # Save to BytesIO
#         file_stream = io.BytesIO()
#         doc.save(file_stream)
#         file_stream.seek(0)
        
#         filename = f"{name.replace(' ', '_')}_Resume.docx" if name else "Resume.docx"
        
#         return StreamingResponse(
#             file_stream,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": f"attachment; filename={filename}"}
#         )

#     except Exception as e:
#         logger.error(f"Error during Word export: {str(e)}")
#         raise HTTPException(status_code=500, detail=str(e))

# @app.get("/health")
# async def health_check():
#     return {"status": "healthy"}





# -------------------


from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import httpx
import os
import json
import logging
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import StreamingResponse
import io
from decouple import config

app = FastAPI()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify the actual origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://host.docker.internal:11434")
# MODEL_NAME = os.getenv("MODEL_NAME", "llama3.2")

OLLAMA_HOST = config("MODEL_HOST_NAME", default="http://host.docker.internal:11434/api/generate")
MODEL_NAME = config("MODEL_NAME", default="llama3.2")
API_KEY = config("API_KEY", default="")

SECTION_LABELS = {
    "summary": "Professional Summary",
    "experience": "Work Experience",
    "education": "Education",
    "projects": "Projects",
    "certifications": "Certifications",
    "skills": "Skills",
    "languages": "Languages",
    "custom": "Additional Information",
}

class ResumeAnalysisRequest(BaseModel):
    resume_data: dict
    job_description: str

class AnalysisResult(BaseModel):
    ats_score: int
    missing_skills: list[str]
    improvements: list[str]


def safe_text(value):
    if value is None:
        return ""
    return str(value).strip()


def extract_json_object(text):
    stripped = safe_text(text)
    if not stripped:
        raise ValueError("Empty model response")

    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{.*\}", stripped, re.DOTALL)
    if not match:
        raise ValueError("Model did not return valid JSON")

    return json.loads(match.group(0))


def sanitize_filename(name):
    cleaned = "".join(ch if ch.isalnum() or ch in (" ", "-", "_") else "_" for ch in name).strip()
    cleaned = "_".join(cleaned.split())
    return cleaned or "Resume"


def format_date_range(start, end, current_label="Present"):
    start = safe_text(start)
    end = safe_text(end)
    if start and end:
        return f"{start} - {end}"
    if start:
        return f"{start} - {current_label}"
    return end


def split_bullet_lines(text):
    if not safe_text(text):
        return []

    lines = []
    for raw_line in str(text).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("•", "-", "*")):
            line = line[1:].strip()
        lines.append(line)
    return lines


def add_paragraph_with_lines(doc, text, style=None):
    lines = split_bullet_lines(text)
    if not lines:
        return

    for line in lines:
        doc.add_paragraph(line, style=style)


def add_bullet_lines(doc, text):
    lines = split_bullet_lines(text)
    if not lines:
        return

    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def configure_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for style_name in ("List Bullet", "List Paragraph"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Calibri"
            doc.styles[style_name].font.size = Pt(10.5)


def add_section_title(doc, text):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(79, 70, 229)


def build_section_order(layout):
    visible_ids = []
    configured_ids = set()
    if isinstance(layout, list):
        for item in layout:
            if not isinstance(item, dict):
                continue
            section_id = safe_text(item.get("id"))
            if section_id:
                configured_ids.add(section_id)
            if item.get("visible", True):
                if section_id:
                    visible_ids.append(section_id)

    for fallback_id in SECTION_LABELS:
        if fallback_id not in visible_ids and fallback_id not in configured_ids:
            visible_ids.append(fallback_id)

    return visible_ids


def has_content(section_id, resume_data):
    personal_info = resume_data.get("personalInfo", {}) or {}

    if section_id == "summary":
        return bool(safe_text(personal_info.get("summary")))
    if section_id in ("experience", "education", "projects", "certifications", "custom"):
        return any(isinstance(item, dict) for item in resume_data.get(section_id if section_id != "custom" else "customSections", []) or [])
    if section_id == "skills":
        return bool(resume_data.get("skills"))
    if section_id == "languages":
        return bool(resume_data.get("languages"))
    return False


def render_resume_doc(doc, resume_data):
    personal_info = resume_data.get("personalInfo", {}) or {}
    full_name = f"{safe_text(personal_info.get('firstName'))} {safe_text(personal_info.get('lastName'))}".strip()

    if full_name:
        heading = doc.add_heading(full_name, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = safe_text(personal_info.get("title"))
    if title:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True

    contact_parts = [
        safe_text(personal_info.get("email")),
        safe_text(personal_info.get("phone")),
        safe_text(personal_info.get("address")),
        safe_text(personal_info.get("linkedin")),
        safe_text(personal_info.get("github")),
        safe_text(personal_info.get("portfolio")),
    ]
    contact_parts = [part for part in contact_parts if part]
    if contact_parts:
        contact_paragraph = doc.add_paragraph(" | ".join(contact_parts))
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = Pt(8)

    for section_id in build_section_order(resume_data.get("layout")):
        if not has_content(section_id, resume_data):
            continue

        add_section_title(doc, SECTION_LABELS.get(section_id, section_id.title()))

        if section_id == "summary":
            add_paragraph_with_lines(doc, personal_info.get("summary"))
            continue

        if section_id == "experience":
            for job in resume_data.get("experience", []) or []:
                if not isinstance(job, dict):
                    continue
                paragraph = doc.add_paragraph()
                role_text = safe_text(job.get("position") or job.get("role"))
                company_text = safe_text(job.get("company"))
                header_text = " | ".join(part for part in [role_text, company_text] if part)
                if header_text:
                    run = paragraph.add_run(header_text)
                    run.bold = True

                date_text = format_date_range(job.get("startDate"), job.get("endDate"))
                if date_text:
                    date_paragraph = doc.add_paragraph()
                    date_run = date_paragraph.add_run(date_text)
                    date_run.italic = True

                address_text = safe_text(job.get("address"))
                if address_text:
                    doc.add_paragraph(address_text)

                add_bullet_lines(doc, job.get("description"))
            continue

        if section_id == "education":
            for edu in resume_data.get("education", []) or []:
                if not isinstance(edu, dict):
                    continue
                paragraph = doc.add_paragraph()
                degree_text = safe_text(edu.get("degree"))
                institution_text = safe_text(edu.get("institution") or edu.get("school"))
                header_text = " | ".join(part for part in [degree_text, institution_text] if part)
                if header_text:
                    run = paragraph.add_run(header_text)
                    run.bold = True

                date_text = format_date_range(edu.get("startDate"), edu.get("endDate"))
                if date_text:
                    date_paragraph = doc.add_paragraph()
                    date_run = date_paragraph.add_run(date_text)
                    date_run.italic = True

                description = safe_text(edu.get("description"))
                if description:
                    add_paragraph_with_lines(doc, description)
            continue

        if section_id == "projects":
            for project in resume_data.get("projects", []) or []:
                if not isinstance(project, dict):
                    continue
                paragraph = doc.add_paragraph()
                title_run = paragraph.add_run(safe_text(project.get("title")))
                title_run.bold = True

                link_text = safe_text(project.get("link"))
                if link_text:
                    link_paragraph = doc.add_paragraph()
                    link_label = link_paragraph.add_run("Link: ")
                    link_label.bold = True
                    link_paragraph.add_run(link_text)

                add_bullet_lines(doc, project.get("description"))
            continue

        if section_id == "certifications":
            for cert in resume_data.get("certifications", []) or []:
                if not isinstance(cert, dict):
                    continue
                paragraph = doc.add_paragraph(style="List Bullet")
                cert_name = safe_text(cert.get("name"))
                cert_issuer = safe_text(cert.get("issuer"))
                cert_date = safe_text(cert.get("date"))
                cert_text = " | ".join(part for part in [cert_name, cert_issuer, cert_date] if part)
                paragraph.add_run(cert_text)
            continue

        if section_id == "skills":
            skills = resume_data.get("skills", []) or []
            skill_list = [safe_text(skill.get("name")) if isinstance(skill, dict) else safe_text(skill) for skill in skills]
            skill_list = [skill for skill in skill_list if skill]
            if skill_list:
                doc.add_paragraph(", ".join(skill_list))
            continue

        if section_id == "languages":
            languages = [safe_text(language) for language in (resume_data.get("languages", []) or [])]
            languages = [language for language in languages if language]
            if languages:
                doc.add_paragraph(", ".join(languages))
            continue

        if section_id == "custom":
            for custom_section in resume_data.get("customSections", []) or []:
                if not isinstance(custom_section, dict):
                    continue
                title_text = safe_text(custom_section.get("title"))
                items = custom_section.get("items", []) or []
                if title_text:
                    paragraph = doc.add_paragraph()
                    title_run = paragraph.add_run(title_text)
                    title_run.bold = True

                for item in items:
                    if not isinstance(item, dict):
                        continue
                    header = safe_text(item.get("header"))
                    sub_header = safe_text(item.get("subHeader"))
                    item_date = safe_text(item.get("date"))
                    description = safe_text(item.get("description"))
                    detail_parts = [part for part in [sub_header, item_date] if part]

                    if header or detail_parts:
                        item_paragraph = doc.add_paragraph(style="List Bullet")
                        if header:
                            header_run = item_paragraph.add_run(header)
                            header_run.bold = True
                        if detail_parts:
                            if header:
                                item_paragraph.add_run(" | ")
                            item_paragraph.add_run(" | ".join(detail_parts))
                        if description:
                            item_paragraph.add_run(f": {description}")
                    elif description:
                        doc.add_paragraph(description, style="List Bullet")


@app.post("/api/analyze-resume-stream")
async def analyze_resume_stream(request: ResumeAnalysisRequest):
    async def event_generator():
        try:
            # 1. Sanitize resume data: Remove heavy base64 'photo' property before sending to LLM
            clean_resume = dict(request.resume_data)
            if clean_resume.get("personalInfo") and "photo" in clean_resume["personalInfo"]:
                clean_resume["personalInfo"] = dict(clean_resume["personalInfo"])
                clean_resume["personalInfo"].pop("photo", None)

            prompt = f"""
            You are an expert ATS (Applicant Tracking System) analyst and career coach.
            Analyze the following resume data against the provided job description.
            
            ### Job Description:
            {request.job_description}
            
            ### Resume Data:
            {json.dumps(clean_resume, indent=2)}
            
            ### Goal:
            1. ATS match score (0-100).
            2. List missing skills from JD.
            3. 3-5 bullet point improvements.
            
            Return ONLY a valid JSON object.
            Do not include markdown fences.
            Do not include any explanation before or after the JSON.

            ### Output Format (JSON only):
            {{
                "ats_score": number,
                "missing_skills": ["skill1", "skill2"],
                "improvements": ["tip1", "tip2"]
            }}
            """

            headers = {
                "Content-Type": "application/json"
            }
            if API_KEY:
                headers["Authorization"] = f"Bearer {API_KEY}"

            async with httpx.AsyncClient(timeout=600.0) as client:
                async with client.stream(
                    "POST",
                    # f"{OLLAMA_HOST}/api/generate",
                    f"{OLLAMA_HOST}",
                    headers=headers,
                    json={
                        "model": MODEL_NAME,
                        "messages": [{ "role": "user", "content": prompt }],
                        "temperature": 0.1,
                        "max_completion_tokens": 4096,
                        "stream": True,
                        "format": "json"
                    }
                ) as response:
                    response.raise_for_status()
                    full_response = ""
                    async for line in response.aiter_lines():
                        if not line:
                            continue
                        try:
                            # 1. Cloud Provider (OpenAI format) Streams
                            if line.startswith("data: "):
                                payload = line[6:].strip()
                                if payload == "[DONE]":
                                    try:
                                        parsed = extract_json_object(full_response)
                                        yield f"data: {json.dumps({'done': True, 'full_response': json.dumps(parsed)})}\n\n"
                                    except Exception as parse_error:
                                        yield f"data: {json.dumps({'error': f'Model did not return valid JSON: {str(parse_error)}', 'full_response': full_response})}\n\n"
                                    break
                                data = json.loads(payload)
                                choices = data.get("choices", [])
                                if choices:
                                    delta = choices[0].get("delta", {})
                                    chunk = delta.get("content", "")
                                    if chunk:
                                        full_response += chunk
                                        yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                            else:
                                # 2. Local Ollama Formats
                                data = json.loads(line)
                                chunk = data.get("response", "")
                                if chunk:
                                    full_response += chunk
                                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                                
                                if data.get("done"):
                                    try:
                                        parsed = extract_json_object(full_response)
                                        yield f"data: {json.dumps({'done': True, 'full_response': json.dumps(parsed)})}\n\n"
                                    except Exception as parse_error:
                                        yield f"data: {json.dumps({'error': f'Model did not return valid JSON: {str(parse_error)}', 'full_response': full_response})}\n\n"
                                    break
                        except json.JSONDecodeError:
                            continue

        except Exception as e:
            logger.error(f"Streaming error: {repr(e)}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.post("/api/export-word")
async def export_word(resume_data: dict):
    try:
        doc = Document()
        configure_document_styles(doc)
        render_resume_doc(doc, resume_data)

        personal_info = resume_data.get("personalInfo", {}) or {}
        full_name = f"{safe_text(personal_info.get('firstName'))} {safe_text(personal_info.get('lastName'))}".strip()
        filename = f"{sanitize_filename(full_name)}_Resume.docx" if full_name else "Resume.docx"

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        logger.error(f"Error during Word export: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    return {"status": "healthy"}
